from docx import Document
from docx.shared import Pt
import os
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 设置对其方式


def txt_to_word(txt_folder, word_file):
    # 创建一个Word文档对象
    doc = Document(word_file)
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    # 获取txt文件夹中的所有txt文件，并按照创建时间排序
    txt_files = sorted(
        [f for f in os.listdir(txt_folder) if f.endswith('.txt')],
        key=lambda f: os.path.getctime(os.path.join(txt_folder, f))
    )

    # 遍历排序后的txt文件列表
    for txt_file in reversed(txt_files[:12]):
        # 使用txt文件名作为章节名
        chapter_title = os.path.splitext(txt_file)[0]

        # 打开txt文件并读取内容
        with open(os.path.join(txt_folder, txt_file), 'r', encoding='utf-8') as txt:
            chapter_content = txt.read()

        # 添加章节标题（居中）和内容到Word文档中

        run = doc.add_heading('', level=2)
        run.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置为左对齐
        run.paragraph_format.space_before = Pt(7)  # 设置段前 0 磅
        run.paragraph_format.space_after = Pt(12)  # 设置段后 0 磅
        heading = run.add_run(chapter_title)
        heading.font.name = u'微软雅黑'
        heading._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置为宋体，和上边的一起使用
        heading.font.size = Pt(20)  # 小四
        heading.alignment = 1  # 设置居中
        heading.bold = True  # 设置为粗体

        # 设置段落内容字体和字号
        paragraph_content = doc.add_paragraph(chapter_content)
        font = paragraph_content.style.font
        font.size = Pt(12)  # 小四

    # 保存Word文档
    doc.save(word_file)
    print(f'Word文档已保存为: {word_file}')


if __name__ == "__main__":
    # 设置txt文件夹路径和要生成的Word文件名
    txt_folder = '诡秘之主'
    word_file = '诡秘之主.docx'

    # 将txt内容按照创建时间写入Word文档
    txt_to_word(txt_folder, word_file)
